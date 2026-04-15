from io import BytesIO
from docx import Document


def build_report_docx_bytes(result: dict) -> bytes:
    doc = Document()

    summary = result.get("summary") or {}
    info = result.get("info") or {}
    visual = result.get("visual") or {}
    audio = result.get("audio") or {}
    rows = result.get("rows") or []
    transcript = result.get("transcript") or ""

    # Title
    doc.add_heading("AI Amendment Bot — QC Report", 0)

    # Scores
    doc.add_heading("Scores", level=1)
    doc.add_paragraph(f"Story Clarity: {summary.get('story_score', 3)}/5")
    doc.add_paragraph(f"Information Clarity: {info.get('score', 3)}/5")
    doc.add_paragraph(f"Visuals: {visual.get('score', 3)}/5")
    doc.add_paragraph(f"Audio: {audio.get('score', 3)}/5")
    doc.add_paragraph(f"Predicted Retention: {summary.get('retention', 'Medium')}")

    # Overall Review
    doc.add_heading("Overall Review", level=1)
    doc.add_paragraph(summary.get("overall_review", ""))

    # Suggestions
    suggestions = summary.get("suggestions", [])
    if suggestions:
        doc.add_heading("Top Suggestions", level=1)
        for s in suggestions:
            doc.add_paragraph(s, style="List Bullet")

    # Information Clarity
    doc.add_heading("Information Clarity", level=1)
    doc.add_paragraph(info.get("summary", ""))

    if info.get("strengths"):
        doc.add_paragraph("Strengths:", style="List Bullet")
        for s in info["strengths"]:
            doc.add_paragraph(s, style="List Bullet 2")

    if info.get("improvements"):
        doc.add_paragraph("Improvements:", style="List Bullet")
        for s in info["improvements"]:
            doc.add_paragraph(s, style="List Bullet 2")

    # Visual Review
    doc.add_heading("Visual Review", level=1)
    doc.add_paragraph(visual.get("summary", ""))

    if visual.get("strengths"):
        doc.add_paragraph("Strengths:", style="List Bullet")
        for s in visual["strengths"]:
            doc.add_paragraph(s, style="List Bullet 2")

    if visual.get("issues"):
        doc.add_paragraph("Issues:", style="List Bullet")
        for s in visual["issues"]:
            doc.add_paragraph(s, style="List Bullet 2")

    if visual.get("suggestions"):
        doc.add_paragraph("Suggestions:", style="List Bullet")
        for s in visual["suggestions"]:
            doc.add_paragraph(s, style="List Bullet 2")

    # Audio Review
    doc.add_heading("Audio Review", level=1)
    doc.add_paragraph(audio.get("summary", ""))

    if audio.get("strengths"):
        doc.add_paragraph("Strengths:", style="List Bullet")
        for s in audio["strengths"]:
            doc.add_paragraph(s, style="List Bullet 2")

    if audio.get("issues"):
        doc.add_paragraph("Issues:", style="List Bullet")
        for s in audio["issues"]:
            doc.add_paragraph(s, style="List Bullet 2")

    if audio.get("suggestions"):
        doc.add_paragraph("Suggestions:", style="List Bullet")
        for s in audio["suggestions"]:
            doc.add_paragraph(s, style="List Bullet 2")

    # QC Table
    if rows:
        doc.add_heading("QC Board", level=1)

        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Type"
        hdr_cells[1].text = "Location"
        hdr_cells[2].text = "Snippet"
        hdr_cells[3].text = "Issue"
        hdr_cells[4].text = "Suggestion"

        for row in rows:
            cells = table.add_row().cells
            cells[0].text = str(row.get("Type", ""))
            cells[1].text = str(row.get("Location", ""))
            cells[2].text = str(row.get("Snippet", ""))
            cells[3].text = str(row.get("Issue", ""))
            cells[4].text = str(row.get("Suggestion", ""))

    # Transcript
    doc.add_heading("Transcript", level=1)
    doc.add_paragraph(transcript)

    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()